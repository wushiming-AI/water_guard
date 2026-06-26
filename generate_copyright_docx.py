#!/usr/bin/env python3
"""
软件著作权申请材料 Word 文档生成器
生成 3 份说明书 .docx + 3 份源代码 .docx
"""
import os
import sys
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

PROJECT_DIR = r"C:\Users\AA\Desktop\water_guard"
OUTPUT_DIR = os.path.join(PROJECT_DIR, "软著申请材料")

# ─── 3 个软著的文件分组 ──────────────────────────────────────────────────
COPYRIGHTS = [
    {
        "name": "水域溺水防控AI预警系统",
        "version": "V4.0",
        "files": [
            "backend.py", "drowning_yolo.py", "pool_map_routes.py", "notify_api.py",
            "init.sql", "generate_demo.py", "generate_icons.py",
            "index.html", "dashboard.html", "alarms.html", "devices.html",
            "users.html", "settings.html", "pool_map.html",
        ],
    },
    {
        "name": "多模态水域安全智能检测系统",
        "version": "V1.0",
        "files": [
            "trackers/__init__.py", "trackers/deep_sort.py", "trackers/trajectory.py",
            "models/__init__.py", "models/cbam.py", "models/cbam_yolo.py", "train_cbam.py",
            "fusion/__init__.py", "fusion/dual_modal.py", "fusion/dual_modal_config.py",
            "fusion/dual_stream.py", "fusion/ir_camera_simulator.py",
            "fusion/ir_preprocess.py", "ir_fusion.py",
        ],
    },
    {
        "name": "水域安全监控移动端管理系统",
        "version": "V1.0",
        "files": [
            "notify/__init__.py", "notify/channels.py", "notify/manager.py", "notify/settings.py",
            "miniprogram/app.js", "miniprogram/app.json", "miniprogram/app.wxss",
            "miniprogram/sitemap.json", "miniprogram/project.config.json",
            "miniprogram/utils/api.js",
            "miniprogram/custom-tab-bar/index.js", "miniprogram/custom-tab-bar/index.json",
            "miniprogram/custom-tab-bar/index.wxml", "miniprogram/custom-tab-bar/index.wxss",
            "miniprogram/pages/alarms/alarms.js", "miniprogram/pages/alarms/alarms.json",
            "miniprogram/pages/alarms/alarms.wxml", "miniprogram/pages/alarms/alarms.wxss",
            "miniprogram/pages/dashboard/dashboard.js", "miniprogram/pages/dashboard/dashboard.json",
            "miniprogram/pages/dashboard/dashboard.wxml", "miniprogram/pages/dashboard/dashboard.wxss",
            "miniprogram/pages/devices/devices.js", "miniprogram/pages/devices/devices.json",
            "miniprogram/pages/devices/devices.wxml", "miniprogram/pages/devices/devices.wxss",
            "miniprogram/pages/login/login.js", "miniprogram/pages/login/login.json",
            "miniprogram/pages/login/login.wxml", "miniprogram/pages/login/login.wxss",
            "miniprogram/pages/settings/settings.js", "miniprogram/pages/settings/settings.json",
            "miniprogram/pages/settings/settings.wxml", "miniprogram/pages/settings/settings.wxss",
        ],
    },
]

LINES_PER_PAGE = 50
MAX_PAGES = 60


# ═══════════════════════════════════════════════════════════════════════
#  辅助函数
# ═══════════════════════════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, color="CCCCCC", size="1"):
    """设置单元格边框"""
    border = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    cell._tc.get_or_add_tcPr().append(border)


def add_table_from_data(doc, headers, rows, header_color="D5E8F0"):
    """从数据列表创建表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        set_cell_shading(cell, header_color)
        set_cell_border(cell)

    # 数据行
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            set_cell_border(cell)

    doc.add_paragraph()  # 表格后空行
    return table


def add_heading_styled(doc, text, level=1, color=RGBColor(0x15, 0x65, 0xC0)):
    """添加带样式的标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = color
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return heading


def add_paragraph_styled(doc, text, bold_prefix=None):
    """添加段落，可带粗体前缀"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(22)
    if bold_prefix:
        run_bold = p.add_run(bold_prefix)
        run_bold.bold = True
        run_bold.font.size = Pt(12)
        run_bold.font.name = "宋体"
        run_bold._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return p


def add_bullet_list(doc, items):
    """添加无序列表"""
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        # 处理加粗前缀
        if isinstance(item, tuple):
            prefix, rest = item
            run_bold = p.add_run(prefix)
            run_bold.bold = True
            run_bold.font.size = Pt(12)
            run_bold.font.name = "宋体"
            run_bold._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run = p.add_run(rest)
        else:
            run = p.add_run(item)
        run.font.size = Pt(12)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_code_block(doc, code_text):
    """添加代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # 设置段落背景
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
    p._element.get_or_add_pPr().append(shading)

    run = p.add_run(code_text)
    run.font.size = Pt(9)
    run.font.name = "Consolas"
    return p


def add_cover_page(doc, title, version, info_lines):
    """添加封面页"""
    # 顶部空白
    for _ in range(6):
        doc.add_paragraph()

    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.bold = True
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 版本
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(version)
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 空行
    for _ in range(4):
        doc.add_paragraph()

    # 信息
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}：{value}")
        run.font.size = Pt(14)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    doc.add_page_break()


def add_toc_page(doc, sections):
    """添加目录页"""
    p = doc.add_paragraph()
    run = p.add_run("目  录")
    run.font.size = Pt(18)
    run.bold = True
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    for section in sections:
        p = doc.add_paragraph()
        run = p.add_run(section)
        run.font.size = Pt(12)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()


def setup_document_styles(doc):
    """设置文档默认样式"""
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


# ═══════════════════════════════════════════════════════════════════════
#  软著1 说明书
# ═══════════════════════════════════════════════════════════════════════

def generate_spec_1(doc):
    """软著1：水域溺水防控AI预警系统 V4.0 说明书"""
    setup_document_styles(doc)

    add_cover_page(doc, "水域溺水防控AI预警系统", "V4.0", [
        ("软件名称", "水域溺水防控AI预警系统"),
        ("软件版本", "V4.0"),
        ("软件类型", "应用软件 / 人工智能监控"),
        ("编程语言", "Python、JavaScript、HTML/CSS、SQL"),
        ("运行平台", "Windows / Linux"),
        ("编写日期", "2026年6月"),
    ])

    add_toc_page(doc, [
        "一、软件概述",
        "二、功能模块说明",
        "三、运行环境",
        "四、安装与部署",
        "五、使用说明",
        "六、技术架构",
        "七、关键技术",
        "八、数据库设计",
        "九、API接口说明",
    ])

    # 一、软件概述
    add_heading_styled(doc, "一、软件概述", level=1)
    add_paragraph_styled(doc, "水域溺水防控AI预警系统是一套基于深度学习的水域安全智能监控平台，通过实时分析视频流中的人员行为，自动识别溺水风险并触发多渠道报警。系统采用 YOLOv8 目标检测算法结合自定义训练模型，能够在泳池、水库、海滩等水域场景下实现 7×24 小时不间断监控。")

    add_heading_styled(doc, "1.1 软件用途", level=2)
    add_paragraph_styled(doc, "本系统主要应用于以下场景：")
    add_bullet_list(doc, [
        "公共泳池安全监控与溺水预警",
        "学校/社区泳池安全辅助管理",
        "水库、河流等开放水域的危险行为监测",
        "水上乐园等娱乐场所安全保障",
    ])

    add_heading_styled(doc, "1.2 软件特点", level=2)
    add_table_from_data(doc, ["特性", "说明"], [
        ["实时AI检测", "基于 YOLOv8 深度学习模型，支持溺水(drowning)、游泳(swimming)、玩耍(playing)三种行为识别"],
        ["多摄像头支持", "支持同时接入多路视频流，统一管理、独立检测、集中展示"],
        ["实时数据推送", "WebSocket 双向通信，前端实时刷新视频帧和报警信息"],
        ["多端联动报警", "支持短信、声光报警器、微信公众号、微信小程序订阅消息、WebSocket 推送五种报警渠道"],
        ["Web管理后台", "提供仪表盘、报警管理、设备管理、用户管理、系统设置、泳池平面图6大功能页面"],
        ["微信小程序", "移动端实时查看监控画面、接收报警推送、管理设备"],
    ])

    # 二、功能模块说明
    add_heading_styled(doc, "二、功能模块说明", level=1)

    add_heading_styled(doc, "2.1 后端服务模块 (backend.py)", level=2)
    add_paragraph_styled(doc, "系统核心后端服务，基于 FastAPI 异步框架构建，提供以下功能：")
    add_bullet_list(doc, [
        ("用户认证：", "手机号+密码登录，JWT令牌管理，角色权限控制（admin/operator）"),
        ("摄像头管理：", "摄像头的注册、删除、心跳检测、帧推送、快照获取、视频流输出"),
        ("报警管理：", "接收YOLO检测模块推送的报警，存入数据库，通过WebSocket推送前端"),
        ("设备管理：", "管理前端设备（摄像头、报警器等）的增删改查"),
        ("系统设置：", "报警灵敏度、推送频率、冷却时间等参数配置"),
        ("数据统计：", "在线摄像头数、今日报警数、累计报警数等实时统计"),
        ("WebSocket服务：", "实时推送视频帧、报警事件、统计数据到前端"),
        ("静态页面托管：", "内置6个HTML管理页面"),
    ])

    add_heading_styled(doc, "2.2 AI检测模块 (drowning_yolo.py)", level=2)
    add_paragraph_styled(doc, "基于 YOLOv8 的实时溺水检测引擎，独立运行于检测主机上：")
    add_bullet_list(doc, [
        ("视频源接入：", "支持 USB摄像头、RTSP网络摄像头、视频文件三种输入"),
        ("行为识别：", "drowning（溺水）、swimming（游泳）、playing（玩耍）三类目标检测"),
        ("多目标跟踪：", "集成 DeepSORT 算法，为每个目标分配唯一ID并绘制运动轨迹"),
        ("报警触发：", "检测到溺水行为且置信度超过阈值时，自动推送报警到后端"),
        ("帧推送：", "以可配置帧率（默认5fps）将带标注的视频帧推送到后端"),
        ("断线重连：", "视频源断开后自动重连，最大重连次数可配置"),
    ])

    add_heading_styled(doc, "2.3 Web前端模块", level=2)
    add_table_from_data(doc, ["页面", "文件", "功能"], [
        ["登录页", "index.html", "用户登录，手机号+密码验证"],
        ["监控仪表盘", "dashboard.html", "多路实时视频监控、报警列表、统计数据"],
        ["报警管理", "alarms.html", "历史报警记录查看、筛选、处理"],
        ["设备管理", "devices.html", "摄像头和设备注册、配置、状态监控"],
        ["用户管理", "users.html", "系统用户增删改、角色分配"],
        ["系统设置", "settings.html", "系统参数配置"],
        ["泳池平面图", "pool_map.html", "SVG动态平面图，实时标注报警位置"],
    ])

    add_heading_styled(doc, "2.4 泳池平面图模块 (pool_map_routes.py + pool_map.html)", level=2)
    add_paragraph_styled(doc, "提供泳池俯视平面的可视化展示：")
    add_bullet_list(doc, [
        "支持矩形、L型、圆形三种泳池形状",
        "SVG动态渲染，实时标注摄像头和报警位置",
        "WebSocket实时更新报警标记",
        "可配置泳池尺寸、泳道数量",
    ])

    # 三、运行环境
    add_heading_styled(doc, "三、运行环境", level=1)

    add_heading_styled(doc, "3.1 硬件环境", level=2)
    add_table_from_data(doc, ["组件", "最低配置", "推荐配置"], [
        ["CPU", "Intel i5-10400 / AMD R5 3600", "Intel i7-12700 / AMD R7 5800X"],
        ["GPU", "NVIDIA GTX 1660 (6GB)", "NVIDIA RTX 3060+ (8GB+)"],
        ["内存", "8 GB", "16 GB+"],
        ["硬盘", "50 GB SSD", "100 GB NVMe SSD"],
        ["网络", "千兆有线", "千兆有线 + 4G备用"],
    ])

    add_heading_styled(doc, "3.2 软件环境", level=2)
    add_table_from_data(doc, ["软件", "版本要求", "用途"], [
        ["Python", "3.9 ~ 3.12", "后端服务 + AI检测"],
        ["MySQL", "5.7+ / 8.0", "数据存储"],
        ["CUDA", "11.8+ / 12.x", "GPU加速推理"],
        ["Node.js", "16+（仅开发需要）", "前端开发工具"],
        ["浏览器", "Chrome 90+ / Edge 90+", "Web管理后台访问"],
    ])

    add_heading_styled(doc, "3.3 依赖库", level=2)
    add_code_block(doc, "fastapi, uvicorn, aiomysql, ultralytics (YOLOv8), opencv-python,\nnumpy<2.0, torch, torchvision, requests, pydantic, aiohttp,\nscipy (DeepSORT依赖), filterpy (卡尔曼滤波)")

    # 四、安装与部署
    add_heading_styled(doc, "四、安装与部署", level=1)

    add_heading_styled(doc, "4.1 数据库初始化", level=2)
    add_code_block(doc, "mysql -u root -p < init.sql")
    add_paragraph_styled(doc, "init.sql 脚本将创建 drowning_alarm 数据库及 users、cameras、devices、alarms、settings 五张数据表。")

    add_heading_styled(doc, "4.2 安装Python依赖", level=2)
    add_code_block(doc, "pip install fastapi uvicorn aiomysql ultralytics opencv-python numpy torch requests pydantic aiohttp")

    add_heading_styled(doc, "4.3 启动后端服务", level=2)
    add_code_block(doc, "python backend.py\n# 服务启动于 http://0.0.0.0:8000")

    add_heading_styled(doc, "4.4 启动AI检测", level=2)
    add_code_block(doc, 'python drowning_yolo.py --camera-id CAM-001 --source 0 --location "泳池A区"\n# --source 0: USB摄像头 / RTSP地址 / 视频文件路径\n# --track True: 启用DeepSORT多目标跟踪')

    add_heading_styled(doc, "4.5 访问Web管理后台", level=2)
    add_paragraph_styled(doc, "浏览器打开 http://127.0.0.1:8000/ 进入登录页，默认账号：13800138000 / 密码：123456")

    # 五、使用说明
    add_heading_styled(doc, "五、使用说明", level=1)

    add_heading_styled(doc, "5.1 用户登录", level=2)
    add_paragraph_styled(doc, "打开系统首页，输入手机号和密码进行登录。系统支持 admin（管理员）和 operator（操作员）两种角色，管理员可管理用户和系统设置。")

    add_heading_styled(doc, "5.2 监控仪表盘", level=2)
    add_paragraph_styled(doc, "登录后进入仪表盘页面，页面分为三个区域：")
    add_bullet_list(doc, [
        ("顶部统计栏：", "显示在线摄像头数、今日报警数、累计报警数、系统运行时间"),
        ("中间视频区：", "以网格布局展示所有在线摄像头的实时画面，点击可全屏查看"),
        ("底部报警栏：", "实时滚动显示最新报警事件"),
    ])

    add_heading_styled(doc, "5.3 报警管理", level=2)
    add_paragraph_styled(doc, "在报警管理页面可以：")
    add_bullet_list(doc, [
        "按时间范围、摄像头、报警类型筛选历史报警记录",
        "查看报警截图和详细信息（置信度、位置、时间）",
        "标记报警处理状态（已处理/误报/待处理）",
    ])

    add_heading_styled(doc, "5.4 摄像头管理", level=2)
    add_paragraph_styled(doc, "在设备管理页面可以添加、编辑、删除摄像头，配置摄像头名称、位置、RTSP地址等信息。系统自动检测摄像头在线状态（30秒无心跳判定为离线）。")

    add_heading_styled(doc, "5.5 泳池平面图", level=2)
    add_paragraph_styled(doc, "泳池平面图页面以SVG方式展示泳池俯视图，实时标注各摄像头位置和报警事件位置，便于快速定位事故区域。")

    add_heading_styled(doc, "5.6 系统设置", level=2)
    add_paragraph_styled(doc, "在设置页面可配置：报警置信度阈值、报警冷却时间、帧推送频率、WebSocket心跳超时等参数。")

    # 六、技术架构
    add_heading_styled(doc, "六、技术架构", level=1)
    add_heading_styled(doc, "6.1 整体架构", level=2)
    add_code_block(doc,
        "┌─────────────────────────────────────────────────────────┐\n"
        "│                    Web 管理后台 (6页)                      │\n"
        "│         index / dashboard / alarms / devices             │\n"
        "│              users / settings / pool_map                 │\n"
        "└──────────────────────┬──────────────────────────────────┘\n"
        "                       │ HTTP + WebSocket\n"
        "┌──────────────────────┴──────────────────────────────────┐\n"
        "│              FastAPI 后端服务 (backend.py)                 │\n"
        "│  ┌─────────┐ ┌─────────┐ ┌─────────┐ ┌───────────────┐  │\n"
        "│  │用户认证  │ │摄像头管理│ │报警管理  │ │ WebSocket推送  │  │\n"
        "│  └─────────┘ └─────────┘ └─────────┘ └───────────────┘  │\n"
        "│  ┌─────────┐ ┌─────────┐ ┌─────────┐ ┌───────────────┐  │\n"
        "│  │设备管理  │ │数据统计  │ │系统设置  │ │ 静态页面托管   │  │\n"
        "│  └─────────┘ └─────────┘ └─────────┘ └───────────────┘  │\n"
        "│  ┌─────────────────┐ ┌─────────────────────────────┐    │\n"
        "│  │ 报警通知管理      │ │ 泳池平面图路由               │    │\n"
        "│  │ (notify_api.py)  │ │ (pool_map_routes.py)       │    │\n"
        "│  └─────────────────┘ └─────────────────────────────┘    │\n"
        "└──────────────────────┬──────────────────────────────────┘\n"
        "                       │ aiomysql\n"
        "┌──────────────────────┴──────────────────────────────────┐\n"
        "│                   MySQL 数据库                            │\n"
        "│    users / cameras / devices / alarms / settings         │\n"
        "└─────────────────────────────────────────────────────────┘\n"
        "\n"
        "┌─────────────────────────────────────────────────────────┐\n"
        "│           YOLO 检测引擎 (drowning_yolo.py)                │\n"
        "│  ┌──────────┐  ┌──────────┐  ┌──────────────────────┐   │\n"
        "│  │视频源接入 │→│ YOLOv8   │→│ DeepSORT多目标跟踪    │   │\n"
        "│  │USB/RTSP  │  │ 行为检测  │  │ 轨迹绘制+风险评分    │   │\n"
        "│  └──────────┘  └──────────┘  └──────────────────────┘   │\n"
        "│                      │ HTTP POST                        │\n"
        "│                      ↓                                  │\n"
        "│              推送帧+报警到后端                            │\n"
        "└─────────────────────────────────────────────────────────┘"
    )

    add_heading_styled(doc, "6.2 数据流", level=2)
    add_paragraph_styled(doc, "系统数据流分为两条主线：")
    add_bullet_list(doc, [
        ("视频流：", "摄像头 → YOLO检测引擎 → 标注帧 → HTTP POST → 后端缓存 → WebSocket → 前端展示"),
        ("报警流：", "YOLO检测到溺水 → HTTP POST /alarm → 后端存库 → WebSocket推送前端 + 通知管理器多渠道报警"),
    ])

    # 七、关键技术
    add_heading_styled(doc, "七、关键技术", level=1)

    add_heading_styled(doc, "7.1 YOLOv8 目标检测", level=2)
    add_paragraph_styled(doc, "系统采用 Ultralytics YOLOv8 模型进行实时目标检测。模型基于自定义标注的泳池溺水数据集训练，能够识别 drowning、swimming、playing 三类行为。检测置信度阈值可配置，默认 0.6。")

    add_heading_styled(doc, "7.2 WebSocket 实时通信", level=2)
    add_paragraph_styled(doc, "前后端通过 WebSocket 建立双向通信通道，实现：")
    add_bullet_list(doc, [
        "视频帧实时推送（二进制JPEG数据）",
        "报警事件即时通知（JSON消息）",
        "统计数据定时刷新",
        "心跳保活机制（60秒超时断开）",
    ])

    add_heading_styled(doc, "7.3 多摄像头并发管理", level=2)
    add_paragraph_styled(doc, "后端采用 asyncio 异步IO模型，每个摄像头的帧缓存和心跳检测独立管理。摄像头帧以 Dict[str, bytes] 存储，配合 asyncio.Event 实现帧更新通知机制，避免轮询开销。")

    add_heading_styled(doc, "7.4 DeepSORT 多目标跟踪", level=2)
    add_paragraph_styled(doc, "集成 DeepSORT 算法（8状态卡尔曼滤波 + 匈牙利算法数据关联），为视频中每个检测目标分配唯一跟踪ID，绘制运动轨迹线，计算风险分数，显著降低误报率。")

    add_heading_styled(doc, "7.5 JWT 令牌认证", level=2)
    add_paragraph_styled(doc, "用户登录后生成 JWT 令牌，前端存储于 localStorage，每次API请求携带 Authorization 头。令牌有效期可配置，过期后需重新登录。")

    # 八、数据库设计
    add_heading_styled(doc, "八、数据库设计", level=1)
    add_table_from_data(doc, ["表名", "说明", "主要字段"], [
        ["users", "系统用户", "id, phone, password, role, nickname, created_at"],
        ["cameras", "摄像头信息", "camera_id, name, location, status, created_at"],
        ["devices", "前端设备", "id, device_name, device_type, camera_id, location, status"],
        ["alarms", "报警记录", "id, camera_id, alarm_type, confidence, image_path, location, created_at"],
        ["settings", "系统设置", "id, key, value, updated_at"],
    ])

    # 九、API接口说明
    add_heading_styled(doc, "九、API接口说明", level=1)
    add_table_from_data(doc, ["接口", "方法", "说明"], [
        ["/api/login", "POST", "用户登录"],
        ["/api/users", "GET/POST/DELETE", "用户管理"],
        ["/api/cameras", "GET/POST/DELETE", "摄像头管理"],
        ["/api/cameras/{id}/frame", "POST", "推送视频帧"],
        ["/api/cameras/{id}/feed", "GET", "获取视频流"],
        ["/api/cameras/{id}/heartbeat", "POST", "摄像头心跳"],
        ["/alarm", "POST", "推送报警"],
        ["/alarms", "GET", "查询报警列表"],
        ["/api/stats", "GET", "获取统计数据"],
        ["/api/devices", "GET/POST/PUT/DELETE", "设备管理"],
        ["/api/settings", "GET/POST", "系统设置"],
        ["/api/notify/*", "GET/POST", "通知管理"],
        ["/api/pool-map/*", "GET/POST", "泳池平面图配置"],
        ["/ws", "WebSocket", "实时通信通道"],
    ])


# ═══════════════════════════════════════════════════════════════════════
#  软著2 说明书
# ═══════════════════════════════════════════════════════════════════════

def generate_spec_2(doc):
    """软著2：多模态水域安全智能检测系统 V1.0 说明书"""
    setup_document_styles(doc)

    add_cover_page(doc, "多模态水域安全智能检测系统", "V1.0", [
        ("软件名称", "多模态水域安全智能检测系统"),
        ("软件版本", "V1.0"),
        ("软件类型", "人工智能算法模块"),
        ("编程语言", "Python"),
        ("运行平台", "Windows / Linux"),
        ("编写日期", "2026年6月"),
    ])

    add_toc_page(doc, [
        "一、软件概述",
        "二、功能模块说明",
        "三、运行环境",
        "四、安装与部署",
        "五、使用说明",
        "六、技术架构",
        "七、关键技术",
        "八、模块接口说明",
    ])

    # 一、软件概述
    add_heading_styled(doc, "一、软件概述", level=1, color=RGBColor(0x2E, 0x7D, 0x32))
    add_paragraph_styled(doc, "多模态水域安全智能检测系统是一套基于深度学习的多模态融合检测算法模块，为水域溺水防控系统提供增强型检测能力。系统包含三大核心算法模块：基于 DeepSORT 的多目标跟踪、基于 CBAM 注意力机制的目标检测增强、以及可见光-红外双模态融合检测。")

    add_heading_styled(doc, "1.1 软件用途", level=2, color=RGBColor(0x2E, 0x7D, 0x32))
    add_bullet_list(doc, [
        "为水域监控场景提供多目标持续跟踪与轨迹分析",
        "通过注意力机制增强 YOLOv8 检测精度，降低漏检率",
        "融合可见光与红外热像数据，实现全天候（含夜间）溺水检测",
        "提供模拟器支持无硬件环境下的算法开发和演示",
    ])

    add_heading_styled(doc, "1.2 软件特点", level=2, color=RGBColor(0x2E, 0x7D, 0x32))
    add_table_from_data(doc, ["特性", "说明"], [
        ["DeepSORT跟踪", "8状态卡尔曼滤波 + 匈牙利算法，稳定跟踪多目标，绘制运动轨迹"],
        ["CBAM注意力", "通道+空间双重注意力机制，注入YOLOv8 backbone提升特征提取"],
        ["双模态融合", "Early/Late/Decision三种融合策略，适配不同硬件配置"],
        ["IR模拟器", "内置红外热像模拟器，无需真实IR硬件即可开发调试"],
        ["风险评分", "基于轨迹特征和行为模式计算实时风险分数"],
    ], header_color="E8F5E9")

    # 二、功能模块说明
    add_heading_styled(doc, "二、功能模块说明", level=1, color=RGBColor(0x2E, 0x7D, 0x32))

    add_heading_styled(doc, "2.1 多目标跟踪模块 (trackers/)", level=2, color=RGBColor(0x2E, 0x7D, 0x32))
    add_paragraph_styled(doc, "基于 DeepSORT 算法的多目标跟踪系统，包含以下组件：")

    add_heading_styled(doc, "2.1.1 DeepSORT 跟踪器 (deep_sort.py)", level=3, color=RGBColor(0x2E, 0x7D, 0x32))
    add_bullet_list(doc, [
        ("Track 类：", "单个跟踪目标的状态管理，包含8维状态向量（x, y, w, h, vx, vy, vw, vh）"),
        ("卡尔曼滤波：", "预测目标下一帧位置，更新观测状态"),
        ("匈牙利算法：", "基于 IoU 的数据关联，匹配检测框与跟踪轨迹"),
        ("生命周期管理：", "Confirmed（确认跟踪）、Tentative（试探期）、Deleted（删除）三态管理"),
        ("参数可配置：", "max_age（跟踪保留帧数）、min_hits（确认命中次数）、iou_threshold（匹配阈值）"),
    ])

    add_heading_styled(doc, "2.1.2 轨迹分析器 (trajectory.py)", level=3, color=RGBColor(0x2E, 0x7D, 0x32))
    add_bullet_list(doc, [
        "记录每个跟踪ID的历史位置序列",
        "绘制渐变色运动轨迹线（越新越亮）",
        "计算运动速度、方向变化率",
        "基于轨迹特征的风险评分算法",
    ])

    add_heading_styled(doc, "2.2 CBAM 注意力模块 (models/)", level=2, color=RGBColor(0x2E, 0x7D, 0x32))
    add_paragraph_styled(doc, "CBAM (Convolutional Block Attention Module) 注意力机制模块，增强 YOLOv8 的特征提取能力：")

    add_heading_styled(doc, "2.2.1 CBAM 模块 (cbam.py)", level=3, color=RGBColor(0x2E, 0x7D, 0x32))
    add_bullet_list(doc, [
        ("ChannelAttention：", "通道注意力子模块，通过全局平均池化和最大池化提取通道权重"),
        ("SpatialAttention：", "空间注意力子模块，通过通道维度的池化提取空间权重"),
        ("CBAM：", "组合通道+空间注意力，先通道后空间，自适应加权特征图"),
        "支持可配置的缩减比率（reduction_ratio），默认16",
    ])

    add_heading_styled(doc, "2.2.2 CBAM-YOLO 集成 (cbam_yolo.py)", level=3, color=RGBColor(0x2E, 0x7D, 0x32))
    add_bullet_list(doc, [
        "将 CBAM 模块注入 YOLOv8 backbone 的 C2f 模块之后",
        "在 Neck 层的 FPN 结构中添加 CBAM 增强特征融合",
        "支持加载预训练权重并保留新增 CBAM 参数",
        "前向传播时自动应用注意力加权",
    ])

    add_heading_styled(doc, "2.2.3 训练脚本 (train_cbam.py)", level=3, color=RGBColor(0x2E, 0x7D, 0x32))
    add_bullet_list(doc, [
        "基于 YOLOv8 训练框架的 CBAM 微调脚本",
        "支持冻结 backbone 层，仅训练 CBAM 参数",
        "可配置学习率、batch size、训练轮数",
        "自动保存最优模型和训练日志",
    ])

    add_heading_styled(doc, "2.3 双模态融合模块 (fusion/)", level=2, color=RGBColor(0x2E, 0x7D, 0x32))
    add_paragraph_styled(doc, "可见光与红外热像的双模态融合检测系统：")

    add_heading_styled(doc, "2.3.1 红外预处理 (ir_preprocess.py)", level=3, color=RGBColor(0x2E, 0x7D, 0x32))
    add_bullet_list(doc, [
        "红外图像增强：对比度拉伸、伪彩映射",
        "温度数据归一化",
        "可见光-红外图像配准（空间对齐）",
        "支持多种红外分辨率输入",
    ])

    add_heading_styled(doc, "2.3.2 双模态融合 (dual_modal.py)", level=3, color=RGBColor(0x2E, 0x7D, 0x32))
    add_bullet_list(doc, [
        ("Early Fusion：", "像素级融合，在输入层合并双模态图像"),
        ("Late Fusion：", "决策级融合，分别检测后合并结果"),
        ("Decision Fusion：", "基于置信度加权的决策融合"),
        "支持配置融合策略和权重参数",
    ])

    add_heading_styled(doc, "2.3.3 双流处理 (dual_stream.py)", level=3, color=RGBColor(0x2E, 0x7D, 0x32))
    add_bullet_list(doc, [
        "可见光流和红外流的同步采集与缓冲",
        "时间戳对齐，确保双模态帧同步",
        "异步处理管道，提高吞吐量",
    ])

    add_heading_styled(doc, "2.3.4 IR 摄像头模拟器 (ir_camera_simulator.py)", level=3, color=RGBColor(0x2E, 0x7D, 0x32))
    add_bullet_list(doc, [
        "从可见光图像模拟生成红外热像",
        "基于亮度-温度映射模型",
        "添加红外噪声和模糊效果",
        "无需真实IR硬件即可开发和演示双模态功能",
    ])

    add_heading_styled(doc, "2.3.5 融合配置 (dual_modal_config.py)", level=3, color=RGBColor(0x2E, 0x7D, 0x32))
    add_bullet_list(doc, [
        "DualModalConfig 配置类，管理融合参数",
        "预设配置：indoor_pool（室内泳池）、outdoor（户外水域）、night（夜间模式）",
        "支持运行时动态调整参数",
    ])

    add_heading_styled(doc, "2.3.6 IR 融合引擎 (ir_fusion.py)", level=3, color=RGBColor(0x2E, 0x7D, 0x32))
    add_bullet_list(doc, [
        "整合双流处理 + 模拟器 + 融合策略的顶层引擎",
        "统一输入输出接口",
        "支持单模态降级运行（IR不可用时自动切换）",
    ])

    # 三、运行环境
    add_heading_styled(doc, "三、运行环境", level=1, color=RGBColor(0x2E, 0x7D, 0x32))
    add_table_from_data(doc, ["组件", "要求"], [
        ["CPU", "Intel i5+ / AMD R5+"],
        ["GPU", "NVIDIA GTX 1660+ (CUDA 11.8+)"],
        ["内存", "8 GB+"],
        ["Python", "3.9 ~ 3.12"],
        ["PyTorch", "2.0+ (with CUDA)"],
        ["Ultralytics", "8.0+ (YOLOv8)"],
        ["SciPy", "1.10+ (匈牙利算法)"],
        ["FilterPy", "1.4+ (卡尔曼滤波)"],
        ["OpenCV", "4.8+"],
    ], header_color="E8F5E9")

    # 四、安装与部署
    add_heading_styled(doc, "四、安装与部署", level=1, color=RGBColor(0x2E, 0x7D, 0x32))
    add_code_block(doc, "pip install torch torchvision ultralytics opencv-python numpy scipy filterpy")
    add_paragraph_styled(doc, "将 trackers/、models/、fusion/ 目录放置在项目根目录下，确保 Python 可导入。")

    # 五、使用说明
    add_heading_styled(doc, "五、使用说明", level=1, color=RGBColor(0x2E, 0x7D, 0x32))

    add_heading_styled(doc, "5.1 启用多目标跟踪", level=2, color=RGBColor(0x2E, 0x7D, 0x32))
    add_code_block(doc, "python drowning_yolo.py --source video.mp4 --track True --trajectory-len 30")
    add_paragraph_styled(doc, "--track 参数控制是否启用 DeepSORT 跟踪，--trajectory-len 设置轨迹显示点数。")

    add_heading_styled(doc, "5.2 使用 CBAM 增强模型", level=2, color=RGBColor(0x2E, 0x7D, 0x32))
    add_code_block(doc, 'from models.cbam_yolo import load_cbam_model\nmodel = load_cbam_model("best_cbam.pt")\nresults = model(frame)')

    add_heading_styled(doc, "5.3 CBAM 模型训练", level=2, color=RGBColor(0x2E, 0x7D, 0x32))
    add_code_block(doc, "python train_cbam.py --data dataset.yaml --epochs 100 --batch 16")

    add_heading_styled(doc, "5.4 双模态融合检测", level=2, color=RGBColor(0x2E, 0x7D, 0x32))
    add_code_block(doc, 'from fusion.ir_fusion import IRFusionEngine\nengine = IRFusionEngine(config="indoor_pool")\nresult = engine.process(visible_frame, ir_frame)')

    add_heading_styled(doc, "5.5 使用 IR 模拟器", level=2, color=RGBColor(0x2E, 0x7D, 0x32))
    add_code_block(doc, "from fusion.ir_camera_simulator import IRCameraSimulator\nsim = IRCameraSimulator()\nir_frame = sim.generate(visible_frame)")

    # 六、技术架构
    add_heading_styled(doc, "六、技术架构", level=1, color=RGBColor(0x2E, 0x7D, 0x32))
    add_code_block(doc,
        "┌──────────────────────────────────────────────────────────┐\n"
        "│                   多模态水域安全智能检测系统                  │\n"
        "├──────────────┬───────────────┬───────────────────────────┤\n"
        "│  DeepSORT    │    CBAM       │    双模态融合              │\n"
        "│  跟踪模块     │   注意力模块   │    fusion/                │\n"
        "├──────────────┼───────────────┼───────────────────────────┤\n"
        "│ trackers/    │ models/       │ fusion/                   │\n"
        "│ ├ deep_sort  │ ├ cbam        │ ├ ir_preprocess           │\n"
        "│ │  .py       │ │  .py        │ ├ dual_modal              │\n"
        "│ ├ trajectory │ ├ cbam_yolo   │ │  .py                    │\n"
        "│ │  .py       │ │  .py        │ ├ dual_stream             │\n"
        "│ └ __init__   │ ├ train_cbam  │ │  .py                    │\n"
        "│              │ │  .py        │ ├ ir_camera_simulator     │\n"
        "│              │ └ __init__    │ │  .py                    │\n"
        "│              │               │ ├ dual_modal_config       │\n"
        "│              │               │ │  .py                    │\n"
        "│              │               │ └ ir_fusion.py            │\n"
        "├──────────────┼───────────────┼───────────────────────────┤\n"
        "│ 卡尔曼滤波    │ 通道+空间     │ Early/Late/Decision       │\n"
        "│ 匈牙利算法    │ 注意力加权     │ 三种融合策略               │\n"
        "│ 轨迹分析      │ YOLOv8集成    │ IR模拟器                  │\n"
        "│ 风险评分      │ 迁移学习      │ 双流同步                   │\n"
        "└──────────────┴───────────────┴───────────────────────────┘"
    )

    # 七、关键技术
    add_heading_styled(doc, "七、关键技术", level=1, color=RGBColor(0x2E, 0x7D, 0x32))

    add_heading_styled(doc, "7.1 DeepSORT 多目标跟踪", level=2, color=RGBColor(0x2E, 0x7D, 0x32))
    add_paragraph_styled(doc, "采用 8 维状态空间的卡尔曼滤波器（x, y, w, h, vx, vy, vw, vh），通过预测-更新循环维持目标状态。使用匈牙利算法基于 IoU 进行检测框与跟踪轨迹的数据关联。跟踪目标生命周期包含 Tentative（试探期，前3帧）、Confirmed（确认跟踪）、Deleted（超时删除）三个状态。")

    add_heading_styled(doc, "7.2 CBAM 注意力机制", level=2, color=RGBColor(0x2E, 0x7D, 0x32))
    add_paragraph_styled(doc, "CBAM 模块包含两个子模块：")
    add_bullet_list(doc, [
        ("通道注意力：", "对特征图进行全局平均池化和最大池化，通过共享MLP网络生成通道权重，使模型关注重要通道"),
        ("空间注意力：", "在通道维度进行平均池化和最大池化，通过7x7卷积生成空间权重，使模型关注重要区域"),
    ])
    add_paragraph_styled(doc, "CBAM 模块嵌入 YOLOv8 的 C2f 模块之后，对 backbone 提取的特征进行自适应加权，增强对水中人员特征的敏感度。")

    add_heading_styled(doc, "7.3 双模态融合策略", level=2, color=RGBColor(0x2E, 0x7D, 0x32))
    add_table_from_data(doc, ["策略", "融合层级", "特点"], [
        ["Early Fusion", "像素级", "在输入层合并双模态图像，信息保留最完整"],
        ["Late Fusion", "决策级", "分别检测后合并结果，计算量低"],
        ["Decision Fusion", "加权决策", "基于置信度加权融合，灵活度高"],
    ], header_color="E8F5E9")

    add_heading_styled(doc, "7.4 IR 模拟器", level=2, color=RGBColor(0x2E, 0x7D, 0x32))
    add_paragraph_styled(doc, "红外摄像头模拟器基于亮度-温度映射模型，从可见光图像推算热分布，添加红外传感器噪声和光学模糊，生成逼真的模拟红外图像。使开发者无需购买红外热像仪即可完成双模态算法开发和演示。")

    add_heading_styled(doc, "7.5 轨迹风险评分", level=2, color=RGBColor(0x2E, 0x7D, 0x32))
    add_paragraph_styled(doc, "基于跟踪轨迹的运动特征（速度变化、方向突变、停留时间）计算实时风险分数。当游泳者突然减速并改变运动方向时，风险分数升高，辅助判断潜在的溺水行为。")

    # 八、模块接口说明
    add_heading_styled(doc, "八、模块接口说明", level=1, color=RGBColor(0x2E, 0x7D, 0x32))
    add_table_from_data(doc, ["模块", "接口", "说明"], [
        ["DeepSORT", "update(detections) → tracks", "输入检测框列表，输出跟踪结果"],
        ["Trajectory", "update(track_id, point) → line", "更新轨迹，返回绘制线段"],
        ["CBAM", "forward(x) → x_weighted", "对特征图进行注意力加权"],
        ["CBAM-YOLO", "load_cbam_model(path) → model", "加载CBAM增强的YOLO模型"],
        ["IR预处理", "preprocess(ir_image) → tensor", "红外图像预处理"],
        ["双模态融合", "fuse(vis, ir) → result", "融合双模态数据"],
        ["双流处理", "sync_grab() → (vis, ir)", "同步采集双模态帧"],
        ["IR模拟器", "generate(vis_frame) → ir_frame", "从可见光生成模拟红外"],
        ["融合引擎", "process(vis, ir) → detection", "端到端融合检测"],
    ], header_color="E8F5E9")


# ═══════════════════════════════════════════════════════════════════════
#  软著3 说明书
# ═══════════════════════════════════════════════════════════════════════

def generate_spec_3(doc):
    """软著3：水域安全监控移动端管理系统 V1.0 说明书"""
    setup_document_styles(doc)

    add_cover_page(doc, "水域安全监控移动端管理系统", "V1.0", [
        ("软件名称", "水域安全监控移动端管理系统"),
        ("软件版本", "V1.0"),
        ("软件类型", "移动应用 / 通知服务"),
        ("编程语言", "JavaScript (微信小程序) + Python"),
        ("运行平台", "微信小程序 / 服务端"),
        ("编写日期", "2026年6月"),
    ])

    add_toc_page(doc, [
        "一、软件概述",
        "二、功能模块说明",
        "三、运行环境",
        "四、安装与部署",
        "五、使用说明",
        "六、技术架构",
        "七、关键技术",
        "八、接口说明",
    ])

    # 一、软件概述
    add_heading_styled(doc, "一、软件概述", level=1, color=RGBColor(0xE6, 0x51, 0x00))
    add_paragraph_styled(doc, "水域安全监控移动端管理系统由两部分组成：微信小程序客户端和多端联动报警通知服务。小程序为管理人员提供移动端的实时监控、报警查看和设备管理能力；通知服务提供短信、声光报警器、微信公众号、小程序订阅消息、WebSocket 五种报警渠道的统一管理和调度。")

    add_heading_styled(doc, "1.1 软件用途", level=2, color=RGBColor(0xE6, 0x51, 0x00))
    add_bullet_list(doc, [
        "管理人员通过手机随时随地查看水域监控画面",
        "实时接收溺水报警推送通知",
        "远程管理摄像头和报警设备",
        "多渠道报警通知确保报警信息必达",
    ])

    add_heading_styled(doc, "1.2 软件特点", level=2, color=RGBColor(0xE6, 0x51, 0x00))
    add_table_from_data(doc, ["特性", "说明"], [
        ["微信小程序", "无需下载安装，扫码即用，5个功能页面+自定义TabBar"],
        ["多渠道报警", "5种报警渠道并行推送，互为备份，确保报警必达"],
        ["渠道管理", "可视化配置各渠道开关、参数、静默时段"],
        ["报警测试", "支持单渠道测试发送，验证渠道可用性"],
        ["订阅消息", "微信小程序订阅消息推送，用户授权后自动送达"],
    ], header_color="FFF3E0")

    # 二、功能模块说明
    add_heading_styled(doc, "二、功能模块说明", level=1, color=RGBColor(0xE6, 0x51, 0x00))

    add_heading_styled(doc, "2.1 微信小程序客户端 (miniprogram/)", level=2, color=RGBColor(0xE6, 0x51, 0x00))

    add_heading_styled(doc, "2.1.1 登录页 (pages/login/)", level=3, color=RGBColor(0xE6, 0x51, 0x00))
    add_bullet_list(doc, [
        "手机号 + 密码登录",
        "JWT 令牌存储于本地 Storage",
        "登录后自动跳转至仪表盘",
        "支持记住手机号功能",
    ])

    add_heading_styled(doc, "2.1.2 仪表盘 (pages/dashboard/)", level=3, color=RGBColor(0xE6, 0x51, 0x00))
    add_bullet_list(doc, [
        "顶部统计卡片：在线摄像头数、今日报警数、累计报警数",
        "实时视频画面列表（JPEG流刷新）",
        "最新报警事件滚动展示",
        "下拉刷新功能",
    ])

    add_heading_styled(doc, "2.1.3 报警管理 (pages/alarms/)", level=3, color=RGBColor(0xE6, 0x51, 0x00))
    add_bullet_list(doc, [
        "历史报警记录列表，支持分页加载",
        "按报警类型筛选（溺水/游泳/玩耍）",
        "报警详情查看（截图、置信度、时间、位置）",
        "报警处理状态标记",
    ])

    add_heading_styled(doc, "2.1.4 设备管理 (pages/devices/)", level=3, color=RGBColor(0xE6, 0x51, 0x00))
    add_bullet_list(doc, [
        "摄像头列表，显示在线/离线状态",
        "添加新摄像头（名称、位置、RTSP地址）",
        "编辑/删除摄像头",
        "设备状态实时更新",
    ])

    add_heading_styled(doc, "2.1.5 系统设置 (pages/settings/)", level=3, color=RGBColor(0xE6, 0x51, 0x00))
    add_bullet_list(doc, [
        "报警灵敏度配置",
        "推送频率设置",
        "修改密码",
        "退出登录",
    ])

    add_heading_styled(doc, "2.1.6 自定义TabBar (custom-tab-bar/)", level=3, color=RGBColor(0xE6, 0x51, 0x00))
    add_bullet_list(doc, [
        "底部导航栏，4个Tab：仪表盘、报警、设备、设置",
        "自定义样式，支持图标和文字",
        "页面间快速切换",
    ])

    add_heading_styled(doc, "2.1.7 API工具 (utils/api.js)", level=3, color=RGBColor(0xE6, 0x51, 0x00))
    add_bullet_list(doc, [
        "封装所有后端 API 调用",
        "统一请求头管理（JWT 令牌）",
        "错误处理和重试机制",
        "基础URL配置",
    ])

    add_heading_styled(doc, "2.2 多端联动报警服务 (notify/)", level=2, color=RGBColor(0xE6, 0x51, 0x00))

    add_heading_styled(doc, "2.2.1 报警渠道 (channels.py)", level=3, color=RGBColor(0xE6, 0x51, 0x00))
    add_paragraph_styled(doc, "定义5种报警渠道，均继承自 NotificationChannel 抽象基类：")
    add_table_from_data(doc, ["渠道", "类名", "说明"], [
        ["短信通知", "SMSChannel", "集成腾讯云SMS API，发送报警短信到管理人员手机"],
        ["声光报警器", "SoundAlarmChannel", "通过HTTP触发现场声光报警器（支持多设备URL列表）"],
        ["微信公众号", "WeChatOAChannel", "通过微信公众号模板消息推送报警通知"],
        ["小程序订阅", "MiniProgramSubscribeChannel", "微信小程序订阅消息推送，用户授权后自动送达"],
        ["WebSocket", "WebSocketChannel", "实时推送到所有已连接的Web前端和小程序"],
    ], header_color="FFF3E0")

    add_heading_styled(doc, "2.2.2 通知管理器 (manager.py)", level=3, color=RGBColor(0xE6, 0x51, 0x00))
    add_bullet_list(doc, [
        ("NotificationManager：", "统一管理所有渠道的注册、配置和调度"),
        ("ChannelConfig：", "渠道配置类，管理启用状态、静默时段、重试参数"),
        ("broadcast()：", "并行调用所有已启用渠道发送报警，返回各渠道发送结果"),
        ("get_channel_status()：", "获取所有渠道当前状态（启用/静默/错误）"),
        ("get_notification_manager()：", "单例模式获取全局通知管理器实例"),
    ])

    add_heading_styled(doc, "2.2.3 通知设置 (settings.py)", level=3, color=RGBColor(0xE6, 0x51, 0x00))
    add_bullet_list(doc, [
        "通知全局配置管理",
        "静默时段配置（如夜间22:00-次日6:00仅推送紧急报警）",
        "各渠道参数配置（API密钥、设备URL、模板ID等）",
    ])

    add_heading_styled(doc, "2.2.4 通知API (notify_api.py)", level=3, color=RGBColor(0xE6, 0x51, 0x00))
    add_bullet_list(doc, [
        "GET /api/notify/channels — 获取所有渠道状态",
        "GET /api/notify/settings — 获取通知设置",
        "POST /api/notify/settings — 更新通知设置",
        "POST /api/notify/test — 发送测试通知到指定渠道",
        "POST /api/notify/silence/{channel} — 静默指定渠道",
        "POST /api/notify/unsilence/{channel} — 取消静默",
    ])

    # 三、运行环境
    add_heading_styled(doc, "三、运行环境", level=1, color=RGBColor(0xE6, 0x51, 0x00))

    add_heading_styled(doc, "3.1 小程序客户端", level=2, color=RGBColor(0xE6, 0x51, 0x00))
    add_table_from_data(doc, ["项目", "要求"], [
        ["微信开发者工具", "最新稳定版"],
        ["基础库版本", "2.16.0+"],
        ["小程序AppID", "需在微信公众平台注册"],
        ["服务器域名", "需在小程序后台配置request合法域名"],
    ], header_color="FFF3E0")

    add_heading_styled(doc, "3.2 通知服务", level=2, color=RGBColor(0xE6, 0x51, 0x00))
    add_table_from_data(doc, ["项目", "要求"], [
        ["Python", "3.9+"],
        ["aiohttp", "3.8+（异步HTTP请求）"],
        ["腾讯云SDK", "短信渠道需要"],
        ["FastAPI", "0.100+（API路由）"],
    ], header_color="FFF3E0")

    # 四、安装与部署
    add_heading_styled(doc, "四、安装与部署", level=1, color=RGBColor(0xE6, 0x51, 0x00))

    add_heading_styled(doc, "4.1 小程序部署", level=2, color=RGBColor(0xE6, 0x51, 0x00))
    add_bullet_list(doc, [
        "打开微信开发者工具，导入 miniprogram/ 目录",
        "在 project.config.json 中配置自己的 AppID",
        "修改 utils/api.js 中的 BASE_URL 为后端服务地址",
        "在微信公众平台 → 开发管理 → 服务器域名中添加后端地址",
        "点击\"预览\"或\"上传\"发布小程序",
    ])

    add_heading_styled(doc, "4.2 通知服务部署", level=2, color=RGBColor(0xE6, 0x51, 0x00))
    add_code_block(doc, "pip install aiohttp")
    add_paragraph_styled(doc, "通知服务已集成在 backend.py 中，通过 notify_api.py 注册路由。启动后端服务时自动加载。")

    # 五、使用说明
    add_heading_styled(doc, "五、使用说明", level=1, color=RGBColor(0xE6, 0x51, 0x00))

    add_heading_styled(doc, "5.1 小程序使用", level=2, color=RGBColor(0xE6, 0x51, 0x00))
    add_bullet_list(doc, [
        "微信扫码打开小程序",
        "输入手机号和密码登录（与Web端账号通用）",
        "在仪表盘查看实时监控画面和统计数据",
        "报警页面查看历史报警记录",
        "设备页面管理摄像头",
        "设置页面配置系统参数",
    ])

    add_heading_styled(doc, "5.2 报警通知管理", level=2, color=RGBColor(0xE6, 0x51, 0x00))
    add_bullet_list(doc, [
        "访问 Web 后台或调用 API 配置各报警渠道参数",
        "发送测试通知验证渠道可用性",
        "配置静默时段（非紧急时段静默部分渠道）",
        "系统检测到溺水时自动并行触发所有已启用渠道",
    ])

    add_heading_styled(doc, "5.3 订阅消息授权", level=2, color=RGBColor(0xE6, 0x51, 0x00))
    add_paragraph_styled(doc, "用户首次使用小程序时，点击\"开启报警通知\"按钮，微信弹出订阅消息授权框，用户同意后即可接收报警推送。")

    # 六、技术架构
    add_heading_styled(doc, "六、技术架构", level=1, color=RGBColor(0xE6, 0x51, 0x00))
    add_code_block(doc,
        "┌──────────────────────────────────────────────────────────┐\n"
        "│                 水域安全监控移动端管理系统                     │\n"
        "├─────────────────────────┬────────────────────────────────┤\n"
        "│    微信小程序客户端        │      多端联动报警服务             │\n"
        "│    miniprogram/          │      notify/                   │\n"
        "├─────────────────────────┼────────────────────────────────┤\n"
        "│ ┌─────────────────────┐ │ ┌────────────────────────────┐ │\n"
        "│ │ pages/login/        │ │ │ channels.py               │ │\n"
        "│ │  登录页              │ │ │  ├ SMSChannel             │ │\n"
        "├─────────────────────┤ │ │  ├ SoundAlarmChannel      │ │\n"
        "│ │ pages/dashboard/    │ │ │  ├ WeChatOAChannel        │ │\n"
        "│ │  仪表盘              │ │ │  ├ MiniProgramSubscribe  │ │\n"
        "├─────────────────────┤ │ │  └ WebSocketChannel      │ │\n"
        "│ │ pages/alarms/       │ ├─┤                            │ │\n"
        "│ │  报警管理            │ │ │ manager.py                │ │\n"
        "├─────────────────────┤ │ │  NotificationManager      │ │\n"
        "│ │ pages/devices/      │ │ │  ChannelConfig            │ │\n"
        "│ │  设备管理            │ │ │                            │ │\n"
        "├─────────────────────┤ │ │ settings.py               │ │\n"
        "│ │ pages/settings/     │ │ │  通知全局配置               │ │\n"
        "│ │  系统设置            │ │ │                            │ │\n"
        "├─────────────────────┤ │ │ notify_api.py             │ │\n"
        "│ │ custom-tab-bar/     │ │ │  REST API 路由             │ │\n"
        "│ │  底部导航            │ │ │                            │ │\n"
        "├─────────────────────┤ │ └────────────┬───────────────┘ │\n"
        "│ │ utils/api.js        │ │              │                 │\n"
        "│ │  API封装             │ │              │                 │\n"
        "└─────────────────────┘ │              │                 │\n"
        "├─────────────────────────┴──────────────┴─────────────────┤\n"
        "│              HTTP / WebSocket → 后端服务                    │\n"
        "└──────────────────────────────────────────────────────────┘"
    )

    # 七、关键技术
    add_heading_styled(doc, "七、关键技术", level=1, color=RGBColor(0xE6, 0x51, 0x00))

    add_heading_styled(doc, "7.1 微信小程序原生框架", level=2, color=RGBColor(0xE6, 0x51, 0x00))
    add_paragraph_styled(doc, "采用微信小程序原生开发框架，使用 WXML（模板）、WXSS（样式）、JS（逻辑）三件套。相比跨平台框架（uni-app/Taro），原生框架性能最优，包体积最小，且能完整使用微信生态能力（订阅消息、扫码等）。")

    add_heading_styled(doc, "7.2 自定义 TabBar", level=2, color=RGBColor(0xE6, 0x51, 0x00))
    add_paragraph_styled(doc, "使用微信小程序自定义 TabBar 组件（custom-tab-bar/），替代默认底部导航栏。支持自定义图标、文字颜色、选中态样式，实现更灵活的UI设计。通过在 app.json 中配置 \"custom\": true 启用。")

    add_heading_styled(doc, "7.3 抽象渠道基类模式", level=2, color=RGBColor(0xE6, 0x51, 0x00))
    add_paragraph_styled(doc, "通知系统采用抽象基类设计模式，NotificationChannel 定义统一的 send() 接口，各渠道子类（SMS/声光/微信/WebSocket）各自实现发送逻辑。NotificationManager 通过统一接口管理所有渠道，新增渠道只需继承基类并实现 send() 方法，符合开闭原则。")

    add_heading_styled(doc, "7.4 异步并行推送", level=2, color=RGBColor(0xE6, 0x51, 0x00))
    add_paragraph_styled(doc, "NotificationManager.broadcast() 使用 asyncio.gather() 并行调用所有已启用渠道的 send() 方法，确保报警延迟最小化。任一渠道发送失败不影响其他渠道，返回结果包含各渠道的发送状态。")

    add_heading_styled(doc, "7.5 静默时段管理", level=2, color=RGBColor(0xE6, 0x51, 0x00))
    add_paragraph_styled(doc, "支持为每个渠道配置静默时段（如22:00-06:00）。静默时段内，非紧急级别的报警不会通过该渠道发送。紧急报警（溺水）无视静默时段，确保安全事件必达。")

    add_heading_styled(doc, "7.6 JWT 令牌统一认证", level=2, color=RGBColor(0xE6, 0x51, 0x00))
    add_paragraph_styled(doc, "小程序与Web端共享同一套 JWT 认证体系。登录成功后获取令牌存储于 wx.setStorageSync()，每次 API 请求通过 header 携带。utils/api.js 统一封装请求拦截，令牌过期自动跳转登录页。")

    # 八、接口说明
    add_heading_styled(doc, "八、接口说明", level=1, color=RGBColor(0xE6, 0x51, 0x00))

    add_heading_styled(doc, "8.1 小程序API (utils/api.js)", level=2, color=RGBColor(0xE6, 0x51, 0x00))
    add_table_from_data(doc, ["函数", "说明"], [
        ["login(phone, password)", "用户登录"],
        ["getStats()", "获取统计数据"],
        ["getCameras()", "获取摄像头列表"],
        ["getAlarms(page, type)", "获取报警列表"],
        ["addCamera(data)", "添加摄像头"],
        ["deleteCamera(id)", "删除摄像头"],
        ["getSettings()", "获取系统设置"],
        ["updateSettings(data)", "更新系统设置"],
    ], header_color="FFF3E0")

    add_heading_styled(doc, "8.2 通知服务API (notify_api.py)", level=2, color=RGBColor(0xE6, 0x51, 0x00))
    add_table_from_data(doc, ["接口", "方法", "说明"], [
        ["/api/notify/channels", "GET", "获取所有渠道状态"],
        ["/api/notify/settings", "GET", "获取通知设置"],
        ["/api/notify/settings", "POST", "更新通知设置"],
        ["/api/notify/test", "POST", "发送测试通知"],
        ["/api/notify/silence/{channel}", "POST", "静默指定渠道"],
        ["/api/notify/unsilence/{channel}", "POST", "取消静默"],
    ], header_color="FFF3E0")


# ═══════════════════════════════════════════════════════════════════════
#  源代码 Word 文档生成
# ═══════════════════════════════════════════════════════════════════════

def generate_source_code_docx(copyright_info, index):
    """生成源代码 Word 文档"""
    doc = Document()
    setup_document_styles(doc)

    # 设置页眉
    section = doc.sections[0]
    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_p.add_run(f"{copyright_info['name']} {copyright_info['version']} 源程序")
    header_run.font.size = Pt(9)
    header_run.font.name = "宋体"
    header_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 设置页脚
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run("第 ")
    footer_run.font.size = Pt(9)
    # 添加页码字段
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    footer_run._element.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    footer_run._element.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    footer_run._element.append(fldChar2)
    footer_run2 = footer_p.add_run(" 页")
    footer_run2.font.size = Pt(9)

    # 读取所有源文件行
    all_lines = []
    for filepath in copyright_info["files"]:
        full_path = os.path.join(PROJECT_DIR, filepath.replace("/", os.sep))
        if not os.path.exists(full_path):
            print(f"  [跳过] 文件不存在: {filepath}")
            continue
        all_lines.append(f"# ===== 文件: {filepath} =====")
        all_lines.append("")
        with open(full_path, "r", encoding="utf-8", errors="replace") as f:
            for line in f:
                line = line.rstrip("\n\r")
                if len(line) > 110:
                    line = line[:110] + " ..."
                all_lines.append(line)
        all_lines.append("")
        all_lines.append("")

    total_lines = len(all_lines)
    total_pages = (total_lines + LINES_PER_PAGE - 1) // LINES_PER_PAGE
    print(f"  总行数: {total_lines}, 总页数: {total_pages}")

    # 确定页范围
    if total_pages <= MAX_PAGES:
        ranges = [(0, total_lines)]
    else:
        first_end = 30 * LINES_PER_PAGE
        last_start = total_lines - 30 * LINES_PER_PAGE
        ranges = [(0, first_end), (last_start, total_lines)]

    for ri, (start, end) in enumerate(ranges):
        if ri > 0:
            # 添加分隔页
            for _ in range(10):
                doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("（以下为源程序最后部分）")
            run.font.size = Pt(14)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            doc.add_page_break()

        for i in range(start, end, LINES_PER_PAGE):
            page_lines = all_lines[i:i + LINES_PER_PAGE]
            for line in page_lines:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.line_spacing = Pt(12)
                safe = line.replace("\t", "    ")
                run = p.add_run(safe)
                run.font.size = Pt(7.5)
                run.font.name = "Consolas"
            # 每页后分页
            if i + LINES_PER_PAGE < end:
                doc.add_page_break()

    output_name = f"软著{index}_{copyright_info['name']}{copyright_info['version']}_源代码.docx"
    output_path = os.path.join(OUTPUT_DIR, output_name)
    doc.save(output_path)
    print(f"  -> {output_path}")
    return output_path


# ═══════════════════════════════════════════════════════════════════════
#  主函数
# ═══════════════════════════════════════════════════════════════════════

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    print(f"输出目录: {OUTPUT_DIR}\n")

    results = []

    # 生成 3 份说明书 Word
    print("=== 生成说明书 Word 文档 ===\n")

    spec_generators = [generate_spec_1, generate_spec_2, generate_spec_3]

    for i, (cr, gen_func) in enumerate(zip(COPYRIGHTS, spec_generators), 1):
        print(f"[说明书 {i}/3] {cr['name']} {cr['version']}")
        doc = Document()
        gen_func(doc)
        output_name = f"软著{i}_{cr['name']}{cr['version']}_说明书.docx"
        output_path = os.path.join(OUTPUT_DIR, output_name)
        doc.save(output_path)
        print(f"  -> {output_path}")
        results.append(output_path)
        print()

    # 生成 3 份源代码 Word
    print("\n=== 生成源代码 Word 文档 ===\n")

    for i, cr in enumerate(COPYRIGHTS, 1):
        print(f"[源代码 {i}/3] {cr['name']} {cr['version']}")
        path = generate_source_code_docx(cr, i)
        results.append(path)
        print()

    print("\n=== 6 份 Word 文档全部生成完成 ===")
    for p in results:
        size = os.path.getsize(p) / 1024
        print(f"  {os.path.basename(p)}  ({size:.0f} KB)")


if __name__ == "__main__":
    main()
